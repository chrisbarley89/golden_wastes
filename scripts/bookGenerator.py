import json
import re
from docx import Document

# Configuration Section
CONFIG = {
    "class_titles": {
        "Bard": "The Bard",
        "Fighter": "The Fighter",
        "Paladin": "The Paladin",
        "Ranger": "The Ranger",
        "Thief": "The Thief",
        "no_class": "Making Moves",
    },
    "type_titles": {
        "Basic": "Basic Moves",
        "Optional": "Optional Moves",
        "Starting": "Starting Moves",
        "Advanced": "Advanced Moves",
        "Expert": "Expert Moves",
    },
    "capitalize_titles": {
        "class": True,  # Capitalize class titles if True
        "type": True,  # Capitalize type titles if True
        "name": True,   # Capitalize name titles if True
    },
    "renaming": {
        "class": "Class Type",  # Rename class section header
        "type": "Move Category",  # Rename type section header
    }
}

# Function to capitalize text conditionally
def capitalize_text(text, should_capitalize):
    return text.upper() if should_capitalize else text

# Function to clean the description by removing "✴" and adding it before "On a"
def clean_description(description):
    description = description.replace("✴", "")
    description = description.replace("On a", "✴ On a")
    return description

# Function to apply rich text formatting to paragraphs
def apply_rich_text_formatting(paragraph, text):
    """Applies bold, italic, and underline formatting to the text based on Markdown-style syntax."""
    # Split text into parts by Markdown-style symbols
    parts = re.split(r'(\*\*.+?\*\*|__.+?__|\*.+?\*|_.+?_)', text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):  # Bold text
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("__") and part.endswith("__"):  # Underline text
            run = paragraph.add_run(part[2:-2])
            run.underline = True
        elif (part.startswith("*") and part.endswith("*")) or (part.startswith("_") and part.endswith("_")):  # Italic text
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:  # Normal text
            paragraph.add_run(part)

# Function to process and categorize moves
def process_moves(moves_data):
    categorized = {}
    
    for move in moves_data:
        move_class = move.get("class", "no_class")
        move_type = move.get("type")
        move_name = move.get("name")
        move_description = move.get("description")
        
        # Ensure class and type keys exist
        if move_class not in categorized:
            categorized[move_class] = {}
        if move_type not in categorized[move_class]:
            categorized[move_class][move_type] = []

        # Clean description text
        cleaned_description = clean_description(move_description)
        
        # Add the move to the categorized dictionary
        categorized[move_class][move_type].append({
            "name": move_name,
            "description": cleaned_description
        })
    
    return categorized

# Function to add moves into docx
def add_moves_to_docx(doc, categorized_moves):
    for move_class, types in categorized_moves.items():
        # Add class heading
        class_title = CONFIG['class_titles'].get(move_class, "Uncategorized")
        class_title = capitalize_text(class_title, CONFIG['capitalize_titles']['class'])
        doc.add_heading(class_title, level=1)

        for move_type, moves in sorted(types.items(), key=lambda x: list(CONFIG['type_titles']).index(x[0])):
            # Add type heading
            type_title = CONFIG['type_titles'].get(move_type, "Unknown Type")
            type_title = capitalize_text(type_title, CONFIG['capitalize_titles']['type'])
            doc.add_heading(type_title, level=2)

            # Sort moves alphabetically by name
            moves_sorted = sorted(moves, key=lambda m: m['name'])

            for move in moves_sorted:
                # Add move name as Heading 3
                move_name = move['name']
                move_name = capitalize_text(move_name, CONFIG['capitalize_titles']['name'])
                doc.add_heading(move_name, level=3)
                # Add move description under the name with rich text formatting
                paragraph = doc.add_paragraph()
                apply_rich_text_formatting(paragraph, move['description'])

# Main function to process the .json and update the .docx
def main():
    # Load JSON file
    with open('./data/moves.json', 'r', encoding="utf-8") as file:
        moves_data = json.load(file)
    
    # Process moves and categorize them
    categorized_moves = process_moves(moves_data)
    
    # Load the existing Word document
    doc = Document('./local/GoldenWastes.docx')
    
    # Add moves data to the document
    add_moves_to_docx(doc, categorized_moves)
    
    # Save the updated document (overwrites the existing file)
    doc.save('./local/GoldenWastes.docx')
    print("Successfully updated the document")  # Debugging message

if __name__ == "__main__":
    main()
